from pathlib import Path
from typing import List
import pandas as pd
from app.agents.base import BaseAgent
from app.agents.types import ProcessingContext, DocumentType

# Optional; agent radi i bez ovih
try:
    import pymupdf4llm  # layout-aware PDF -> Markdown
except Exception:
    pymupdf4llm = None

try:
    import fitz  # PyMuPDF
except Exception:
    fitz = None

try:
    import pdfplumber  # tabele iz PDF-a
except Exception:
    pdfplumber = None

try:
    import PyPDF2  # fallback tekst iz PDF-a
except Exception:
    PyPDF2 = None


class TextExtractAgent(BaseAgent):
    def __init__(self):
        super().__init__("TextExtractAgent")
    
    async def process(self, context: ProcessingContext) -> ProcessingContext:
        file_path = Path(context.file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {context.file_path}")
        
        # reset per-run
        context.tables = []
        context.images = []

        if context.document_type == DocumentType.PDF:
            self._extract_from_pdf(file_path, context)
        elif context.document_type == DocumentType.DOCX:
            self._extract_from_docx_with_tables(file_path, context)
        elif context.document_type in [DocumentType.CSV, DocumentType.XLSX]:
            context.text_content = self._extract_from_table(file_path, context)
        elif context.document_type == DocumentType.TEXT:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                context.text_content = f.read()
        
        context.metadata['extracted_text_length'] = len(context.text_content or "")
        return context
    
    # ---------- PDF ----------
    def _extract_from_pdf(self, file_path: Path, context: ProcessingContext) -> None:
        assets_dir = file_path.parent / (file_path.stem + "_extract") / "assets"
        assets_dir.mkdir(parents=True, exist_ok=True)
        
        md_text = ""
        used_md = False
        images: List[str] = []
        tables_csv: List[str] = []
        
        # A) Markdown iz PDF-a (ako postoji)
        if pymupdf4llm is not None:
            try:
                md_text = pymupdf4llm.to_markdown(str(file_path), image_path=str(assets_dir))
                used_md = True
            except Exception as e:
                context.metadata.setdefault('warnings', []).append(f"pymupdf4llm failed: {e}")
        
        # B) prikupljene slike (ako ih ima)
        if assets_dir.exists():
            for p in sorted(assets_dir.glob('*')):
                if p.is_file():
                    images.append(str(p))
        
        # C) tabele u CSV (pdfplumber)
        if pdfplumber is not None:
            try:
                with pdfplumber.open(str(file_path)) as pdf:
                    for pno, page in enumerate(pdf.pages, start=1):
                        for ti, table in enumerate(page.extract_tables() or []):
                            if not table:
                                continue
                            headers = table[0] if table[0] else None
                            data = table[1:] if headers else table
                            try:
                                df = pd.DataFrame(data, columns=headers if headers else None)
                            except Exception:
                                max_len = max(len(r) for r in data) if data else 0
                                df = pd.DataFrame(data, columns=[f"col_{i}" for i in range(max_len)])
                            csv_path = assets_dir / f"table_{pno}_{ti}.csv"
                            df.to_csv(csv_path, index=False)
                            tables_csv.append(str(csv_path))
                            preview_text = df.head(20).to_string(index=False)
                            context.tables.append({
                                'page': pno,
                                'csv_path': str(csv_path),
                                'rows': int(len(df)),
                                'cols': int(len(df.columns)),
                                'preview': preview_text,
                            })
            except Exception as e:
                context.metadata.setdefault('warnings', []).append(f"pdfplumber failed: {e}")
        
        # D) Fallback tekst (PyPDF2)
        if not md_text.strip() and PyPDF2 is not None:
            try:
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    parts = []
                    for page in reader.pages:
                        try:
                            t = page.extract_text() or ""
                        except Exception:
                            t = ""
                        if t.strip():
                            parts.append(t.strip())
                    md_text = "\n\n".join(parts)
            except Exception as e:
                context.metadata.setdefault('warnings', []).append(f"PyPDF2 fallback failed: {e}")
        
        context.text_content = md_text.strip()
        for img in images:
            context.images.append({'path': img})
        context.metadata.setdefault('pdf', {})
        context.metadata['pdf']['assets_dir'] = str(assets_dir)
        context.metadata['pdf']['tables_csv'] = tables_csv
        context.metadata['pdf']['used_md'] = used_md
    
    # ---------- DOCX ----------
    def _extract_from_docx_with_tables(self, file_path: Path, context: ProcessingContext) -> None:
        try:
            from docx import Document as DocxDocument
        except Exception as e:
            raise Exception(f"DOCX extraction error (python-docx missing?): {e}")

        assets_dir = file_path.parent / (file_path.stem + "_extract") / "assets"
        assets_dir.mkdir(parents=True, exist_ok=True)

        try:
            doc = DocxDocument(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text]
            context.text_content = ("\n".join(paragraphs)).strip()

            tables_csv: List[str] = []
            for ti, table in enumerate(doc.tables):
                rows = []
                max_cols = 0
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells]
                    rows.append(cells)
                    max_cols = max(max_cols, len(cells))
                norm = [r + [""] * (max_cols - len(r)) for r in rows]
                headers = norm[0] if (norm and all(v != "" for v in norm[0])) else None
                data = norm[1:] if headers else norm
                df = pd.DataFrame(data, columns=headers if headers else None)
                csv_path = assets_dir / f"docx_table_{ti}.csv"
                df.to_csv(csv_path, index=False)
                tables_csv.append(str(csv_path))
                preview = df.head(20).to_string(index=False)
                context.tables.append({
                    "page": None,
                    "csv_path": str(csv_path),
                    "rows": int(len(df)),
                    "cols": int(len(df.columns)),
                    "preview": preview,
                })
            context.metadata.setdefault("docx", {})
            context.metadata["docx"]["tables_csv"] = tables_csv
            context.metadata["docx"]["assets_dir"] = str(assets_dir)
        except Exception as e:
            raise Exception(f"DOCX extraction error: {str(e)}")
    
    # ---------- CSV/XLSX ----------
    def _extract_from_table(self, file_path: Path, context: ProcessingContext) -> str:
        """
        Robusna obrada CSV/XLSX:
        - CSV: auto separator (engine='python'), fallback encoding, on_bad_lines='skip'
        - XLSX: svi sheetovi -> CSV u .../_extract/assets/
        - U context.tables ide preview; text_content dobija kratki spojeni preview
        """
        def _df_preview(df: pd.DataFrame, rows: int = 50) -> str:
            df2 = df.copy().replace({pd.NA: "", None: ""})
            return df2.head(rows).to_string(index=False)

        assets_dir = file_path.parent / (file_path.stem + "_extract") / "assets"
        assets_dir.mkdir(parents=True, exist_ok=True)

        previews: List[str] = []
        suffix = file_path.suffix.lower()
        try:
            if suffix == ".csv":
                try:
                    df = pd.read_csv(file_path, sep=None, engine="python", on_bad_lines="skip")
                except UnicodeDecodeError:
                    df = pd.read_csv(file_path, sep=None, engine="python", on_bad_lines="skip", encoding="latin-1")
                except Exception:
                    df = pd.read_csv(file_path, on_bad_lines="skip")
                df.columns = [str(c).strip() for c in df.columns]
                df = df.replace({pd.NA: "", None: ""})
                preview = _df_preview(df, rows=50)
                previews.append(preview)
                context.tables.append({
                    'page': None,
                    'csv_path': str(file_path),
                    'rows': int(len(df)),
                    'cols': int(len(df.columns)),
                    'preview': preview,
                })
                return df.to_string(index=False)
            else:
                xl = pd.ExcelFile(file_path)
                for sheet_name in xl.sheet_names:
                    try:
                        sdf = xl.parse(sheet_name=sheet_name, dtype=str)
                    except Exception:
                        sdf = xl.parse(sheet_name=sheet_name)
                    sdf.columns = [str(c).strip() for c in sdf.columns]
                    sdf = sdf.replace({pd.NA: "", None: ""})
                    safe_sheet = "".join(ch if ch.isalnum() or ch in ("-","_") else "_" for ch in sheet_name)[:40]
                    csv_path = assets_dir / f"xlsx_{safe_sheet}.csv"
                    try:
                        sdf.to_csv(csv_path, index=False)
                    except Exception:
                        sdf.to_csv(csv_path, index=False, encoding="utf-8", errors="ignore")
                    preview = _df_preview(sdf, rows=50)
                    previews.append(f"[SHEET: {sheet_name}]\n{preview}")
                    context.tables.append({
                        'page': None,
                        'csv_path': str(csv_path),
                        'rows': int(len(sdf)),
                        'cols': int(len(sdf.columns)),
                        'preview': preview,
                    })
                return ("\n\n".join(previews)).strip()
        except Exception as e:
            raise Exception(f"Table extraction error: {str(e)}")
